from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

items = [
    ("1",  "ZD550 Karbon Fiber Frame", 1,
     6858.00, "https://www.f1depo.com/urun/zd550-katlanabilir-inis-takimli-karbon-fiber-drone-govdesi-demonte",
     7162.80, "https://www.robotzade.com/urun/zd550-katlanabilir-inis-takimli-karbon-fiber-drone-govdesi"),

    ("2",  "SunnySky X4110S 400KV Fırçasız Motor", 4,
     2487.24, "https://www.f1depo.com/urun/sunnysky-xs-high-power-x4110s-400kv-fircasiz-motor",
     3158.40, "https://www.f1depo.com/urun/sunnysky-xs-high-power-x4110s-400kv-fircasiz-motor"),

    ("3",  "Hobbywing XRotor 40A ESC", 4,
     836.80,  "https://www.f1depo.com/urun/hobbywing-xrotor-40a-2-6s-esc-2",
     1297.20, "https://www.f1depo.com/urun/hobbywing-xrotor-40a-2-6s-esc-2"),

    ("4",  "1355 Karbon Fiber Pervane Seti (CW/CCW)", 2,
     767.34,  "https://www.robocombo.com/1355-karbon-fiber-drone-pervane-seti-cw-ccw-siyah-3814",
     564.00,  "https://www.f1depo.com/urun/13-inc-3k-karbon-fiber-pervane-1355-2-adet-cw-ccw"),

    ("5",  "Waveshare IMX219-160 Kamera", 1,
     3361.00, "https://www.f1depo.com/urun/imx219-jetson-nano-kamera-nvidia-jetson-ve-raspberry-pi-uyumlu",
     1381.45, "https://www.direnc.net/imx219-kamera-modulu-160-derece-fov-en"),

    ("6",  "NVIDIA Jetson Orin Nano Super Dev Kit", 1,
     20255.26,"https://market.samm.com/en/nvidia-jetson-orin-nano-super-developer-kit",
     19683.60,"https://openzeka.com/urun/nvidia-jetson-orin-nano-developer-kit/"),

    ("7",  "Holybro M9N GPS Modülü", 1,
     3040.00, "https://drone.net.tr/drone-tuning/holybro-holybro-m9n-gps-jst-gh-10-pin-connector-for-gps1-port-on-flight-controller.html",
     3648.00, "https://drone.net.tr/en/drone-tuning/holybro-holybro-m9n-gps-for-pix32-2-4-6.html"),

    ("8",  "Benewake TFmini-S Lidar", 1,
     3427.26, "https://www.robotsepeti.com/benewake-tfs20-l-hassas-lidar-lazer-mesafe-sensoru",
     2904.00, "https://drone.net.tr/en/gimbal-ve-faydali-yukler/tfmini-s-12m-lidar-ranging-module.html"),

    ("9",  "22.2V 6S 7000mAh LiPo Batarya", 1,
     7899.66, "https://www.motorobit.com/222v-6s-7000mah-45c-lipo-batarya",
     8460.00, "https://www.f1depo.com/urun/22-2v-7000mah-40c-lipo-batarya-6s-jetfire-pil"),

    ("10", "Matek PDB-HEX 2-12S Güç Dağıtım Kartı", 1,
     1615.93, "https://www.voltaj.net/matek-pdb-hex-2-12s-5a-5-12v-w-264a-bec-drone-guc-dagitim-karti-pmu35549",
     1733.00, "https://www.komponentci.net/matek-pdb-hex-2-12s-5a-5-12v-w-264a-bec-drone-guc-dagitim-karti-pmu11893"),

    ("11", "Matek UBEC DUO 4A/5V & 4A/12V", 1,
     1700.00, "https://www.karbondrone.com/mateksys-ubec-duo-4a512v--4a5v",
     1700.00, "https://www.thkmodelucak.com/urun/matek-ubec-duo-4a-5-12v-4a-5v-2070"),

    ("12", "u-blox NEO-M8N GPS Modülü", 1,
     584.10,  "https://www.robocombo.com/ublox-gy-gpsv3-neo-8m-m8n-gps-modulu--3196",
     700.23,  "https://www.komponentci.net/ublox-gy-gpsv3-neo-8m-m8n-gps-modulu-pmu11767"),

    ("13", "Holybro SiK Telemetry Radio V3 433MHz", 1,
     6072.17, "https://www.kompent.com/holybro-sik-telemetri-radyo-modulu-v3-100mw-433mhz",
     9300.00, "https://www.aykuthavacilik.com/urun/sik-telemetri-radyo-v3-100mw-433mhz"),

    ("14", "Ebyte LoRa E32-433T20D RF Modülü", 1,
     438.08,  "https://www.robocombo.com/ebyte-e32-433t20d-lora-uart-rf-module-433mhz-20db-4680",
     564.00,  "https://www.f1depo.com/urun/ebyte-e32-433t20d-lora-modul"),

    ("15", "TTGO T-Display ESP32", 1,
     1077.60, "https://www.robolinkmarket.com/ttgo-esp32-ch340k-wifi-bluetooth-modul-gelistirme-karti",
     886.49,  "https://www.hiber.com.tr/ttgo-t-display-esp32-ch340k"),

    ("16", "Waveshare OV5640 Kamera Modülü", 1,
     1612.13, "https://www.f1depo.com/urun/waveshare-ov5640-kamera-modul-2592x1944-balik-gozu",
     1392.99, "https://market.samm.com/en/ov5640-kamera-karti-b-5mp-2592x1944-balikgozu-lens-1"),
]

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

title = doc.add_paragraph()
tr = title.add_run("Malzeme Alım Listesi")
tr.bold = True
tr.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run(
    "Aşağıda projemizde kullanılacak olan malzemelerin güncel fiyat listesi yer almaktadır. "
    "Eski fiyatlar Mayıs 2026'da hazırlanan ilk bütçe raporundaki tutarlardır, yeni fiyatlar "
    "Haziran 2026 itibarıyla ilgili tedarikçilerden alınmıştır."
)

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'

hdr = table.rows[0].cells
labels = ["No", "Malzeme", "Adet", "Eski Fiyat", "Eski Link", "Yeni Fiyat", "Toplam", "Yeni Link"]
for i, h in enumerate(labels):
    hdr[i].text = ""
    run = hdr[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)

def fmt(n):
    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def add_link(cell, url, text):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    r.font.underline = True
    r.font.size = Pt(10)
    part = doc.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    r_elm = r._r
    p._p.remove(r_elm)
    hyperlink.append(r_elm)
    p._p.append(hyperlink)

old_total = 0.0
new_total = 0.0

for no, name, qty, old_unit, old_link, new_unit, new_link in items:
    new_sum = new_unit * qty
    old_total += old_unit * qty
    new_total += new_sum
    row = table.add_row().cells
    row[0].text = no
    row[1].text = name
    row[2].text = str(qty)
    row[3].text = fmt(old_unit)
    add_link(row[4], old_link, "link")
    row[5].text = fmt(new_unit)
    row[6].text = fmt(new_sum)
    add_link(row[7], new_link, "link")
    for c in row:
        for para in c.paragraphs:
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(10)

trow = table.add_row().cells
trow[0].merge(trow[5])
trow[0].text = ""
p = trow[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Genel Toplam").bold = True
trow[6].text = fmt(new_total)
trow[7].text = ""
for c in trow:
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True

widths = [Cm(0.8), Cm(4.6), Cm(0.9), Cm(1.7), Cm(1.2), Cm(1.7), Cm(1.9), Cm(1.2)]
for row in table.rows:
    for i, c in enumerate(row.cells):
        if i < len(widths):
            c.width = widths[i]

doc.add_paragraph()
n = doc.add_paragraph()
n.add_run("Not: ").bold = True
n.add_run(
    "Listeyi hazırlarken her ürünün stok durumu tedarikçi sayfalarından kontrol "
    "edilmiştir. Bazı kalemler için piyasada geçici stok kesintisi bulunmakta olup "
    "bu ürünler için ilgili mağaza ile telefon yoluyla sipariş kaydı oluşturulacaktır. "
    "Holybro M9N GPS ve Benewake TFmini-S drone.net.tr üzerinden ön sipariş ile "
    "temin edilebilmekte olup tedarikçinin belirttiği tahmini teslim süresi "
    "Temmuz - Ağustos 2026'dır. "
    "Fiyatlar KDV dahildir; kargo bedelleri ayrıca hesaplanacaktır."
)

p2 = doc.add_paragraph()
p2.add_run("Eski toplam: ").bold = True
p2.add_run(f"{fmt(old_total)} TL")
p2.add_run("    Yeni toplam: ").bold = True
p2.add_run(f"{fmt(new_total)} TL")

out = "/home/ati/Attia-Pro/Projectos/Teknofest-enes-group/last report/Kokpit_Guncel_Alim_Listesi.docx"
doc.save(out)
print("OK")
print(f"Old: {fmt(old_total)} New: {fmt(new_total)} Diff: {fmt(new_total - old_total)}")
